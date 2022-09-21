import os
from turtle import width
import pandas as pd
from numpy import random
import matplotlib.pyplot as plt
from docx import Document
import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


PINGLIST = []
with open("list.txt") as file:  # list.txt裡是要ping的ip或電腦名稱
    pinglist = file.read()
    pinglist = pinglist.splitlines()
    i = 0
    for i in range(0, len(pinglist)):
        a = pinglist[i].strip()
        if a not in PINGLIST:
            PINGLIST.append(a)
        i += 1

UP = []
DOWN = []
for j in range(0, len(PINGLIST)):
    RESULT = os.system("ping -n 2 " + PINGLIST[j])
    if RESULT == 0:
        UP.append(PINGLIST[j])
    else:
        DOWN.append(PINGLIST[j])
    j += 1

# # to Excel
# # df = pd.DataFrame([UP, DOWN], index=('UP', 'DOWN'))
# # Transdf = df.T
# # Transdf.to_excel('result.xlsx')

# to txt
# 先整理list中前方加入\n斷行
# nUP = []
# u = 0
# for u in range(0, len(UP)):
#     nUP.append("\n"+UP[u])
#     u += 1

# d = 0
# nDOWN = []
# for d in range(0, len(DOWN)):
#     nDOWN.append("\n"+DOWN[d])
#     d += 1

# 算個比例吧，希望能弄出個chart
UPratio = len(UP)/len(PINGLIST)*100
DOWNratio = len(DOWN)/len(PINGLIST)*100
stat = ["%2d%% is up, %2d%% is down" % (UPratio, DOWNratio)]

# # 寫入記事本
# f = open("result.txt", "w+")
# ALL = ["【UP】"] + nUP + ["\n\n【DOWN】"] + nDOWN + ["\n\n"] + stat
# f.writelines(ALL)
# f.close()

# 畫圖啦

labels = 'UP', 'DOWN'
sizes = [UPratio, DOWNratio]
explode = (0, 0.12)  # only "explode" the 2nd slice (i.e. 'Hogs')

fig1, ax1 = plt.subplots()
ax1.pie(sizes, explode=explode, labels=labels, autopct='%1.1f%%',
        shadow=True, startangle=90)
ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
plt.savefig("RatioPie.png")

# 寫入Word，含圖和表格，成為報表
document = Document()
document.add_heading('PING日報表', 0)

loc_dt = datetime.datetime.today()  # local time
loc_dt_format = loc_dt.strftime("%Y/%m/%d %H:%M:%S")
TimeLine = document.add_paragraph(str(loc_dt_format))
TimeLine.alignment = WD_ALIGN_PARAGRAPH.RIGHT
Picture = document.add_picture('./RatioPie.png', width=Inches(4))
PicturePara = document.paragraphs[2]
PicturePara.alignment = WD_ALIGN_PARAGRAPH.CENTER

if len(UP) >= len(DOWN):
    MaxRow = len(UP)
else:
    MaxRow = len(DOWN)

# 創建表格
table = document.add_table(rows=1, cols=2)
# 首列名稱
row = table.rows[0].cells
row[0].text = 'UP'
row[1].text = 'DOWN'

datatable = document.add_table(rows=MaxRow, cols=2)
data_UP = datatable.columns[0].cells
data_DOWN = datatable.columns[1].cells
for i in range(0, len(UP)):
    data_UP[i].text = UP[i]
    i += 1

for j in range(0, len(DOWN)):
    data_DOWN[j].text = DOWN[j]
    j += 1


statLine = document.add_heading(stat, 1)
statLine.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Adding style to a table
table.style = 'Colorful List'
datatable.style = 'Colorful List'

# document.add_page_break()
document.save('PING_report.docx')
